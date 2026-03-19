from __future__ import annotations

from io import BytesIO
from typing import Any, Callable, Dict, List

from docx import Document
from docx.shared import Inches, Pt

Token = Dict[str, Any]


class DocxRenderer:
    _BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    _NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]

    def __init__(
        self,
        doc: Document,
        plantuml_fetcher: Callable[[str], bytes],
        image_width_inches: float = 5.0,
    ) -> None:
        self.doc = doc
        self.plantuml_fetcher = plantuml_fetcher
        self.image_width = Inches(image_width_inches)

    def render(self, tokens: List[Token]) -> None:
        for token in tokens:
            self._render_block(token)

    # ------------------------------------------------------------------
    # Block-level rendering
    # ------------------------------------------------------------------

    def _render_block(self, token: Token) -> None:
        t = token["type"]

        if t == "heading":
            level = min(token["attrs"]["level"], 9)
            p = self._add_paragraph(f"Heading {level}")
            self._render_inline(p, token.get("children", []))

        elif t == "paragraph":
            p = self._add_paragraph("Normal")
            self._render_inline(p, token.get("children", []))

        elif t == "block_code":
            info = (token["attrs"].get("info") or "").strip().lower()
            code = token.get("raw", "")
            if info == "plantuml":
                self._render_plantuml(code)
            else:
                self._render_code_block(code)

        elif t == "list":
            ordered = token["attrs"].get("ordered", False)
            depth = token["attrs"].get("depth", 0)
            for item in token.get("children", []):
                self._render_list_item(item, depth, ordered)

        elif t == "table":
            self._render_table(token)

        elif t == "block_quote":
            for child in token.get("children", []):
                self._render_block(child)

        # blank_line, thematic_break, unknown — skip

    def _render_list_item(self, token: Token, depth: int, ordered: bool) -> None:
        styles = self._NUMBER_STYLES if ordered else self._BULLET_STYLES
        style = styles[min(depth, len(styles) - 1)]

        first_para = True
        for child in token.get("children", []):
            ct = child["type"]
            if ct in ("paragraph", "block_text"):
                if first_para:
                    p = self._add_paragraph(style)
                    self._render_inline(p, child.get("children", []))
                    first_para = False
                else:
                    p = self._add_paragraph("Normal")
                    self._render_inline(p, child.get("children", []))
            elif ct == "list":
                self._render_block(child)
            else:
                self._render_block(child)

    def _render_plantuml(self, code: str) -> None:
        try:
            image_bytes = self.plantuml_fetcher(code)
            self.doc.add_picture(BytesIO(image_bytes), width=self.image_width)
        except Exception as exc:
            p = self._add_paragraph("Normal")
            run = p.add_run(f"[PlantUML error: {exc}]")
            run.italic = True

    def _render_code_block(self, code: str) -> None:
        p = self._add_paragraph("Normal")
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def _render_table(self, token: Token) -> None:
        children = token.get("children", [])
        head = next((c for c in children if c["type"] == "table_head"), None)
        body = next((c for c in children if c["type"] == "table_body"), None)

        head_cells = head.get("children", []) if head else []
        body_rows = body.get("children", []) if body else []

        num_cols = len(head_cells)
        if num_cols == 0:
            return

        table = self.doc.add_table(rows=1 + len(body_rows), cols=num_cols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        # Header row
        for col_idx, cell_token in enumerate(head_cells):
            cell = table.rows[0].cells[col_idx]
            self._render_inline(cell.paragraphs[0], cell_token.get("children", []), bold=True)

        # Body rows
        for row_idx, row_token in enumerate(body_rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_token in enumerate(row_token.get("children", [])):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    self._render_inline(cell.paragraphs[0], cell_token.get("children", []))

    # ------------------------------------------------------------------
    # Inline rendering
    # ------------------------------------------------------------------

    def _render_inline(
        self,
        paragraph: Any,
        children: List[Token],
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        for child in children:
            ct = child["type"]

            if ct == "text":
                run = paragraph.add_run(child.get("raw", ""))
                run.bold = bold or None
                run.italic = italic or None

            elif ct == "strong":
                self._render_inline(paragraph, child.get("children", []), bold=True, italic=italic)

            elif ct == "emphasis":
                self._render_inline(paragraph, child.get("children", []), bold=bold, italic=True)

            elif ct == "codespan":
                run = paragraph.add_run(child.get("raw", ""))
                run.font.name = "Courier New"
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

            elif ct == "softline":
                paragraph.add_run(" ")

            elif ct == "linebreak":
                paragraph.add_run("\n")

            elif ct == "link":
                # Render link text without a hyperlink (not in scope)
                self._render_inline(paragraph, child.get("children", []), bold=bold, italic=italic)

            # raw_html, image, and other types are silently skipped

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _add_paragraph(self, style: str) -> Any:
        try:
            return self.doc.add_paragraph(style=style)
        except KeyError:
            return self.doc.add_paragraph(style="Normal")
